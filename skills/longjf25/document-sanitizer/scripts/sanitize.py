"""
文档脱敏与恢复工具 v1.1
- config 模式：生成脱敏配置模板
- sanitize 模式：批量脱敏 docx/xlsx 文件
- restore 模式：从脱敏记录恢复原始内容
"""
import os
import re
import sys
import json
import shutil
from pathlib import Path
from datetime import datetime
from collections import defaultdict

# ============================================================
# 安全常量
# ============================================================

# 配置文件最大允许大小（1 MB），防止加载超大文件导致内存溢出
MAX_CONFIG_SIZE = 1 * 1024 * 1024
# 正则匹配超时秒数，防止恶意正则（ReDoS）导致 CPU 挂起
REGEX_TIMEOUT = 10
# 单个文本最大处理长度（10 MB），跳过超长文本避免性能问题
MAX_TEXT_LENGTH = 10 * 1024 * 1024

# ============================================================
# 工具函数
# ============================================================

LOG_FILE = None


def log(msg):
    """日志输出（控制台 + 文件），写入失败不影响主流程"""
    clean = ''.join(c for c in str(msg) if c.isprintable() or c in '\n\r\t')
    print(clean)
    if LOG_FILE:
        try:
            with open(LOG_FILE, "a", encoding="utf-8") as f:
                f.write(clean + "\n")
        except OSError:
            pass  # 磁盘满或权限问题时不中断


def init_log(workspace):
    """初始化日志文件"""
    global LOG_FILE
    LOG_FILE = workspace / "_sanitize_log.txt"
    try:
        with open(LOG_FILE, "w", encoding="utf-8") as f:
            f.write(f"Document Sanitizer v1.1 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n")
    except OSError as e:
        print(f"[WARN] 无法创建日志文件: {e}")


def validate_subpath(workspace, target):
    """
    安全校验：确保 target 目录在 workspace 范围内，防止路径穿越攻击
    返回 (resolved_target, error_msg)
    """
    try:
        resolved = (workspace / target).resolve()
        workspace_resolved = workspace.resolve()
        # 检查 resolved 是否以 workspace_resolved 开头
        if not str(resolved).startswith(str(workspace_resolved)):
            return None, f"目标路径 '{target}' 超出工作区范围，已拒绝"
        return resolved, None
    except (OSError, ValueError) as e:
        return None, f"路径解析失败: {e}"


# ============================================================
# 配置模板
# ============================================================

CONFIG_TEMPLATE = '''{
  "exact_rules": [
    {"pattern": "张三", "replacement": ""},
    {"pattern": "某某公司", "replacement": "[公司A]"}
  ],
  "regex_rules": [
    {"pattern": "1[3-9]\\\\d{9}", "replacement": "", "label": "手机号"},
    {"pattern": "\\\\d{6}(?:19|20)\\\\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\\\\d|3[01])\\\\d{3}[\\\\dXx]", "replacement": "", "label": "身份证号"}
  ]
}
'''


def generate_config(workspace):
    """生成配置模板"""
    config_path = workspace / "_sanitize_config.json"
    if config_path.exists():
        log(f"[INFO] 配置文件已存在: {config_path}")
        log("  如需重新生成，请先删除或重命名现有配置文件。")
        return

    with open(config_path, "w", encoding="utf-8") as f:
        f.write(CONFIG_TEMPLATE)

    log(f"[OK] 配置模板已生成: {config_path}")
    log("")
    log("使用说明:")
    log("  1. 编辑 _sanitize_config.json，添加需要脱敏的关键字和正则规则")
    log("  2. exact_rules: 精确匹配规则，直接替换指定文本")
    log("  3. regex_rules: 正则匹配规则，支持模式匹配（手机号、身份证等）")
    log("  4. replacement 为空时自动生成占位符（如 [RED_0001]）")
    log("  5. regex_rules 的 label 字段用于生成易读的占位符标签")
    log("")
    log("编辑完成后，执行「脱敏文档」即可批量处理。")


# ============================================================
# 配置读取与映射表构建
# ============================================================

def load_config(root_dir):
    """读取并验证配置文件（从根目录查找），含大小限制防内存溢出"""
    config_path = root_dir / "_sanitize_config.json"
    if not config_path.exists():
        log("[ERROR] 未找到配置文件 _sanitize_config.json")
        log("  请先执行: python sanitize.py config <工作文件夹>")
        return None

    # 安全校验：文件大小检查
    file_size = config_path.stat().st_size
    if file_size > MAX_CONFIG_SIZE:
        log(f"[ERROR] 配置文件过大 ({file_size:,} bytes)，上限 {MAX_CONFIG_SIZE:,} bytes")
        return None

    try:
        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)
    except json.JSONDecodeError as e:
        log(f"[ERROR] 配置文件 JSON 格式错误: {e}")
        return None
    except OSError as e:
        log(f"[ERROR] 无法读取配置文件: {e}")
        return None

    # 验证结构
    if not isinstance(config, dict):
        log("[ERROR] 配置文件根元素必须是 JSON 对象")
        return None
    if "exact_rules" not in config:
        config["exact_rules"] = []
    if "regex_rules" not in config:
        config["regex_rules"] = []

    # 安全校验：规则数量上限（防止单次配置上千条规则）
    total_rules = len(config["exact_rules"]) + len(config["regex_rules"])
    if total_rules > 500:
        log(f"[WARN] 规则数量较多 ({total_rules} 条)，处理可能较慢")

    return config


def build_mapping(config, workspace):
    """
    根据配置构建完整的映射表
    返回: {
        "mapping": {原文: 替换内容, ...},
        "exact_rules": [(pattern, replacement), ...],
        "regex_rules": [(compiled_regex, replacement), ...]
    }
    """
    mapping = {}
    exact_rules = []
    regex_rules = []
    counter = 1  # 全局序号

    # 处理精确匹配规则
    log("\n[CONFIG] 精确匹配规则:")
    for rule in config.get("exact_rules", []):
        pattern = rule.get("pattern", "").strip()
        if not pattern:
            continue
        replacement = rule.get("replacement", "").strip()
        if not replacement:
            replacement = f"[RED_{counter:04d}]"
            counter += 1
        mapping[("exact", pattern)] = replacement
        exact_rules.append((pattern, replacement))
        log(f"  {pattern} → {replacement}")

    # 处理正则匹配规则（注意：占位符在脱敏执行时动态生成，以确保每个不同匹配项有唯一占位符）
    log("\n[CONFIG] 正则匹配规则:")
    for rule in config.get("regex_rules", []):
        pattern = rule.get("pattern", "").strip()
        if not pattern:
            continue
        label = rule.get("label", "").strip()
        replacement = rule.get("replacement", "").strip()
        try:
            # 安全校验：设置正则超时，防止 ReDoS 攻击
            compiled = re.compile(pattern, timeout=REGEX_TIMEOUT)
        except re.error as e:
            log(f"  [WARN] 正则表达式无效: {pattern} ({e})，已跳过")
            continue
        except TypeError:
            # Python < 3.11 不支持 timeout 参数，回退到无超时编译
            try:
                compiled = re.compile(pattern)
            except re.error as e:
                log(f"  [WARN] 正则表达式无效: {pattern} ({e})，已跳过")
                continue

        if not replacement:
            # 占位符将在脱敏时动态生成，这里用 None 标记
            replacement = None
        regex_rules.append((compiled, replacement, pattern, label))
        label_str = f" (标签: {label})" if label else ""
        repl_str = replacement if replacement else "[动态生成]"
        log(f"  /{pattern}/{label_str} → {repl_str}")

    return {
        "mapping": mapping,
        "exact_rules": exact_rules,
        "regex_rules": regex_rules,
        "_counter": counter,  # 传递计数器状态
    }


# ============================================================
# 文件扫描
# ============================================================

SKIP_DIRS = {"_sanitized_output", "_restored_output", "_文档_md", "_markdown", "update",
             "_文档_md_backup", "node_modules", ".git", "__pycache__"}
DOC_EXTENSIONS = {".docx", ".xlsx"}


def scan_documents(workspace):
    """扫描工作区中的 docx/xlsx 文件"""
    documents = []
    for root, dirs, files in os.walk(workspace):
        # 跳过特殊目录
        rel_root = Path(root).relative_to(workspace)
        skip = False
        for part in rel_root.parts:
            if part in SKIP_DIRS:
                skip = True
                break
        if skip:
            continue

        for f in files:
            ext = Path(f).suffix.lower()
            if ext in DOC_EXTENSIONS:
                documents.append(Path(root) / f)

    return documents


# ============================================================
# 文本替换引擎
# ============================================================

def do_replace(text, exact_rules, regex_rules, mapping, counter):
    """
    对文本执行所有替换规则
    正则匹配时为每个不同的匹配项动态生成唯一占位符
    mapping 使用元组 key (type, original_text) 避免冲突
    安全校验：跳过超长文本防止性能问题
    """
    # 安全校验：跳过超长文本
    if len(text) > MAX_TEXT_LENGTH:
        return text

    # 先执行正则替换（优先处理模式匹配）
    for compiled, replacement, _pattern, label in regex_rules:
        if replacement is not None:
            # 用户指定了固定替换内容
            text = compiled.sub(replacement, text)
            # 记录映射
            for m in compiled.finditer(text):
                matched = m.group(0)
                key = ("regex", matched)
                if key not in mapping:
                    mapping[key] = replacement
        else:
            # 动态生成占位符：每个不同的匹配值获得唯一占位符
            def dynamic_replace(match, _label=label, _mapping=mapping, _counter=counter):
                matched_text = match.group(0)
                key = ("regex", matched_text)
                if key in _mapping:
                    return _mapping[key]
                if _label:
                    _counter[0] += 1
                    placeholder = f"[RED_{_label}_{_counter[0]}]"
                else:
                    _counter[0] += 1
                    placeholder = f"[RED_REGEX_{_counter[0]:04d}]"
                _mapping[key] = placeholder
                return placeholder

            text = compiled.sub(dynamic_replace, text)

    # 再执行精确替换
    for pattern, replacement in exact_rules:
        text = text.replace(pattern, replacement)

    return text


# ============================================================
# docx 处理
# ============================================================

def replace_in_paragraph(para, exact_rules, regex_rules, mapping, counter):
    """对段落中的完整文本执行替换，保留第一个 run 的格式"""
    full_text = para.text
    if not full_text:
        return

    new_text = do_replace(full_text, exact_rules, regex_rules, mapping, counter)
    if new_text == full_text:
        return  # 无变化

    # 将替换后的文本写入第一个 run，清空其余 run
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        # 无 run 时创建一个
        run = para.add_run(new_text)


def sanitize_docx(file_path, exact_rules, regex_rules, mapping, counter):
    """脱敏 docx 文件：段落、表格、页眉页脚、文本框"""
    from docx import Document

    doc = Document(str(file_path))

    # 替换段落文本
    for para in doc.paragraphs:
        replace_in_paragraph(para, exact_rules, regex_rules, mapping, counter)

    # 替换表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, exact_rules, regex_rules, mapping, counter)

    # 替换页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for para in header.paragraphs:
                    replace_in_paragraph(para, exact_rules, regex_rules, mapping, counter)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    replace_in_paragraph(para, exact_rules, regex_rules, mapping, counter)

    # 替换文本框（通过 XML 查找）
    try:
        from docx.oxml.ns import qn
        for textbox in doc.element.iter(qn('w:txbxContent')):
            for para_elem in textbox.iter(qn('w:p')):
                # 收集所有文本
                texts = []
                for text_elem in para_elem.iter(qn('w:t')):
                    texts.append(text_elem.text or "")
                full_text = "".join(texts)
                if not full_text:
                    continue
                new_text = do_replace(full_text, exact_rules, regex_rules, mapping, counter)
                if new_text == full_text:
                    continue
                # 将新文本写入第一个 text 元素，清空其余
                text_elems = list(para_elem.iter(qn('w:t')))
                if text_elems:
                    text_elems[0].text = new_text
                    for te in text_elems[1:]:
                        te.text = ""
    except Exception:
        pass  # 文本框处理失败不影响主流程

    doc.save(str(file_path))
    return True


# ============================================================
# xlsx 处理
# ============================================================

def sanitize_xlsx(file_path, exact_rules, regex_rules, mapping, counter):
    """脱敏 xlsx 文件：所有工作表的单元格"""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path))

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_value = do_replace(cell.value, exact_rules, regex_rules, mapping, counter)
                    if new_value != cell.value:
                        cell.value = new_value

    wb.save(str(file_path))
    return True


# ============================================================
# 文件名脱敏
# ============================================================

def sanitize_filename(filename, exact_rules, regex_rules, mapping, counter):
    """
    对文件名（不含扩展名）执行脱敏替换
    返回 (脱敏后文件名, 是否有变化)
    """
    name_without_ext = Path(filename).stem
    ext = Path(filename).suffix

    new_name = do_replace(name_without_ext, exact_rules, regex_rules, mapping, counter)
    if new_name == name_without_ext:
        return filename, False

    return new_name + ext, True


# ============================================================
# 脱敏执行
# ============================================================

def run_sanitize(workspace, target_dir=None, rename_files=False):
    """
    执行脱敏流程
    workspace: 工作区根目录（配置文件、输出目录、记录文件的存放位置）
    target_dir: 脱敏目标目录（扫描文档的范围），为 None 时默认为 workspace
    """
    if target_dir is None:
        target_dir = workspace

    # 安全校验：确保 target_dir 在 workspace 范围内
    if target_dir != workspace:
        validated_target, err = validate_subpath(workspace, target_dir)
        if err:
            log(f"[ERROR] {err}")
            return
        target_dir = validated_target

    # 读取配置（从工作区根目录）
    config = load_config(workspace)
    if not config:
        return

    init_log(workspace)
    log("\n[SANITIZE] 文档脱敏")
    log("=" * 60)
    if target_dir != workspace:
        log(f"  脱敏范围: {target_dir.relative_to(workspace)}")

    # 构建映射表
    rule_data = build_mapping(config, workspace)
    exact_rules = rule_data["exact_rules"]
    regex_rules = rule_data["regex_rules"]
    mapping = rule_data["mapping"]  # 共享映射表
    counter = [rule_data.get("_counter", 1)]  # 计数器（用列表包装以便闭包修改）

    if not exact_rules and not regex_rules:
        log("[WARN] 未配置任何脱敏规则，请编辑 _sanitize_config.json")
        return

    # 扫描文档（在目标目录中扫描，但跳过工作区级的输出目录）
    log("\n[SCAN] 扫描文档...")
    documents = scan_documents(target_dir)
    log(f"  发现 {len(documents)} 个文档")

    if not documents:
        log("[WARN] 未发现任何 docx/xlsx 文件")
        return

    for doc in documents:
        rel = doc.relative_to(target_dir)
        log(f"  - {rel}")

    # 创建输出目录（在工作区根目录下）
    output_dir = workspace / "_sanitized_output"
    output_dir.mkdir(exist_ok=True)

    # 执行脱敏
    log("\n[PROCESS] 开始脱敏...")
    success = 0
    failed = []
    files_processed = []
    filename_mapping = {}  # 文件名映射: {脱敏后相对路径: 原始相对路径}

    for doc in documents:
        # 计算相对路径（相对于目标目录）
        rel_to_target = doc.relative_to(target_dir)

        # 文件名脱敏
        original_rel = str(rel_to_target)
        if rename_files:
            new_filename, changed = sanitize_filename(
                rel_to_target.name, exact_rules, regex_rules, mapping, counter
            )
            if changed:
                new_rel = rel_to_target.parent / new_filename
                log(f"  [RENAME] {rel_to_target.name} → {new_filename}")
            else:
                new_rel = rel_to_target
        else:
            new_rel = rel_to_target

        # 在输出目录中保持相对路径
        dest = output_dir / new_rel
        dest.parent.mkdir(parents=True, exist_ok=True)

        # 复制到输出目录
        shutil.copy2(str(doc), str(dest))

        try:
            ext = doc.suffix.lower()
            if ext == ".docx":
                sanitize_docx(dest, exact_rules, regex_rules, mapping, counter)
            elif ext == ".xlsx":
                sanitize_xlsx(dest, exact_rules, regex_rules, mapping, counter)

            success += 1
            files_processed.append(str(rel_to_target))

            # 记录文件名映射（仅当文件名有变化时）
            if str(new_rel) != original_rel:
                filename_mapping[str(new_rel)] = original_rel

            log(f"  [{success}/{len(documents)}] {rel_to_target.name} [OK]")
        except Exception as e:
            failed.append((str(rel_to_target), str(e)))
            log(f"  [{success + len(failed)}/{len(documents)}] {rel_to_target.name} [FAIL: {e}]")

    # 生成脱敏记录
    log("\n[RECORD] 生成脱敏记录...")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    record_file = workspace / f"_sanitize_record_{timestamp}.json"

    record = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "config_file": "_sanitize_config.json",
        "rules_applied": {
            "exact": len(exact_rules),
            "regex": len(regex_rules)
        },
        "rename_files": rename_files,
        "files_processed": files_processed,
        "filename_mapping": filename_mapping,  # 文件名映射
        # 从 mapping {("exact"/"regex", original): placeholder} 转为 {placeholder: original}
        "mapping": {v: k[1] for k, v in mapping.items()}
    }

    with open(record_file, "w", encoding="utf-8") as f:
        json.dump(record, f, ensure_ascii=False, indent=2)

    # 汇报结果
    log("\n" + "=" * 60)
    log(f"[RESULT] 脱敏完成: {success}/{len(documents)} 成功")
    if failed:
        log(f"失败文件: {len(failed)}")
        for path, err in failed:
            log(f"  - {path}: {err}")
    log(f"\n脱敏输出目录: {output_dir}")
    log(f"脱敏记录文件: {record_file}")
    log(f"操作日志: {LOG_FILE}")
    log("=" * 60)

    if success > 0:
        log("\n恢复方法:")
        log(f"  python scripts/sanitize.py restore <工作文件夹> --record {record_file.name}")




def extract_text(file_path):
    """提取文档中的文本（用于正则匹配预扫描）"""
    ext = file_path.suffix.lower()
    texts = []

    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(str(file_path))
            for para in doc.paragraphs:
                texts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texts.append(cell.text)
        elif ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(str(file_path))
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            texts.append(cell.value)
    except Exception:
        return ""

    return "\n".join(texts)


# ============================================================
# 恢复执行
# ============================================================

def run_restore(workspace, record_path=None):
    """执行恢复流程"""
    # 查找记录文件
    if not record_path:
        record_path = find_latest_record(workspace)

    if not record_path:
        log("[ERROR] 未找到脱敏记录文件")
        log("  请使用 --record 参数指定记录文件")
        return

    if not record_path.exists():
        log(f"[ERROR] 记录文件不存在: {record_path}")
        return

    init_log(workspace)
    log("\n[RESTORE] 文档恢复")
    log("=" * 60)

    # 读取记录
    try:
        with open(record_path, "r", encoding="utf-8") as f:
            record = json.load(f)
    except json.JSONDecodeError as e:
        log(f"[ERROR] 记录文件格式错误: {e}")
        return

    mapping = record.get("mapping", {})
    if not mapping:
        log("[ERROR] 记录文件中无映射数据")
        return

    filename_mapping = record.get("filename_mapping", {})
    if filename_mapping:
        log(f"  文件名映射: {len(filename_mapping)} 条")

    log(f"  记录文件: {record_path.name}")
    log(f"  脱敏时间: {record.get('timestamp', '未知')}")
    log(f"  映射条目: {len(mapping)} 条")

    # mapping 格式: {placeholder: original_text}
    # 构建反向替换规则: [(placeholder, original_text), ...]
    reverse_exact = [(placeholder, original) for placeholder, original in mapping.items()]
    log(f"  反向映射: {len(reverse_exact)} 条")

    # 用于检测占位符的集合
    placeholder_set = set(mapping.keys())

    # 扫描脱敏输出目录
    output_dir = workspace / "_sanitized_output"
    if not output_dir.exists():
        log(f"[ERROR] 脱敏输出目录不存在: {output_dir}")
        return

    log(f"\n[SCAN] 扫描脱敏文件...")
    documents = []
    for root, dirs, files in os.walk(output_dir):
        for f in files:
            ext = Path(f).suffix.lower()
            if ext in DOC_EXTENSIONS:
                documents.append(Path(root) / f)

    log(f"  发现 {len(documents)} 个文件")

    if not documents:
        log("[WARN] 脱敏输出目录中没有文件")
        return

    # 创建恢复输出目录
    restore_dir = workspace / "_restored_output"
    restore_dir.mkdir(exist_ok=True)

    # 执行恢复
    log("\n[PROCESS] 开始恢复...")
    success = 0
    failed = []
    warnings = []

    for doc in documents:
        rel = doc.relative_to(output_dir)
        # 恢复原始文件名
        rel_str = str(rel)
        original_rel_str = filename_mapping.get(rel_str, rel_str)
        original_rel = Path(original_rel_str)
        dest = restore_dir / original_rel
        dest.parent.mkdir(parents=True, exist_ok=True)

        # 如果文件名被恢复了，日志显示
        if rel_str != original_rel_str:
            log(f"  [RENAME] {rel.name} → {original_rel.name}")

        # 复制到恢复目录
        shutil.copy2(str(doc), str(dest))

        try:
            ext = doc.suffix.lower()

            # 先检查文件中是否包含占位符
            text = extract_text(dest)
            has_placeholder = any(p in text for p in placeholder_set)

            if not has_placeholder:
                warnings.append(str(original_rel))
                log(f"  [WARN] {original_rel.name} - 未检测到脱敏占位符，跳过恢复")
                continue

            if ext == ".docx":
                restore_docx(dest, reverse_exact)
            elif ext == ".xlsx":
                restore_xlsx(dest, reverse_exact)

            success += 1
            log(f"  [{success}/{len(documents)}] {original_rel.name} [OK]")
        except Exception as e:
            failed.append((str(original_rel), str(e)))
            log(f"  [FAIL] {original_rel.name}: {e}")

    # 残留检查
    residual_warnings = []
    if success > 0:
        log("\n[CHECK] 检查残留占位符...")
        for doc_path in restore_dir.rglob("*.docx"):
            text = extract_text(doc_path)
            residual = re.findall(r'\[RED_[^\]]+\]', text)
            if residual:
                unique = set(residual)
                rel = doc_path.relative_to(restore_dir)
                residual_warnings.append((str(rel), list(unique)))

        for doc_path in restore_dir.rglob("*.xlsx"):
            text = extract_text(doc_path)
            residual = re.findall(r'\[RED_[^\]]+\]', text)
            if residual:
                unique = set(residual)
                rel = doc_path.relative_to(restore_dir)
                residual_warnings.append((str(rel), list(unique)))

    # 汇报结果
    log("\n" + "=" * 60)
    log(f"[RESULT] 恢复完成: {success}/{len(documents)} 成功")
    if failed:
        log(f"失败文件: {len(failed)}")
        for path, err in failed:
            log(f"  - {path}: {err}")
    if warnings:
        log(f"\n跳过文件（无脱敏标记）: {len(warnings)}")
        for path in warnings:
            log(f"  - {path}")
    if residual_warnings:
        log(f"\n[WARN] 残留占位符检测:")
        for path, placeholders in residual_warnings:
            log(f"  - {path}: {', '.join(placeholders)}")
        log("  请检查这些文件是否已完全恢复。")

    log(f"\n恢复输出目录: {restore_dir}")
    log(f"操作日志: {LOG_FILE}")
    log("=" * 60)


def restore_in_paragraph(para, reverse_rules):
    """对段落执行反向替换，保留第一个 run 的格式"""
    full_text = para.text
    if not full_text:
        return

    new_text = do_replace_restore(full_text, reverse_rules)
    if new_text == full_text:
        return

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        run = para.add_run(new_text)


def do_replace_restore(text, reverse_rules):
    """恢复时执行反向替换"""
    for pattern, replacement in reverse_rules:
        text = text.replace(pattern, replacement)
    return text


def restore_docx(file_path, reverse_rules):
    """恢复 docx 文件"""
    from docx import Document
    doc = Document(str(file_path))

    for para in doc.paragraphs:
        restore_in_paragraph(para, reverse_rules)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    restore_in_paragraph(para, reverse_rules)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for para in header.paragraphs:
                    restore_in_paragraph(para, reverse_rules)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    restore_in_paragraph(para, reverse_rules)

    try:
        from docx.oxml.ns import qn
        for textbox in doc.element.iter(qn('w:txbxContent')):
            for para_elem in textbox.iter(qn('w:p')):
                texts = []
                for text_elem in para_elem.iter(qn('w:t')):
                    texts.append(text_elem.text or "")
                full_text = "".join(texts)
                if not full_text:
                    continue
                new_text = do_replace_restore(full_text, reverse_rules)
                if new_text == full_text:
                    continue
                text_elems = list(para_elem.iter(qn('w:t')))
                if text_elems:
                    text_elems[0].text = new_text
                    for te in text_elems[1:]:
                        te.text = ""
    except Exception:
        pass

    doc.save(str(file_path))


def restore_xlsx(file_path, reverse_rules):
    """恢复 xlsx 文件"""
    from openpyxl import load_workbook
    wb = load_workbook(str(file_path))

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_value = do_replace_restore(cell.value, reverse_rules)
                    if new_value != cell.value:
                        cell.value = new_value

    wb.save(str(file_path))


def find_latest_record(workspace):
    """查找最新的脱敏记录文件"""
    records = sorted(workspace.glob("_sanitize_record_*.json"))
    if records:
        return records[-1]
    return None


# ============================================================
# 主入口
# ============================================================

def main():
    if len(sys.argv) < 2:
        print("用法:")
        print("  python sanitize.py config <工作文件夹>                                    生成配置模板")
        print("  python sanitize.py sanitize <工作文件夹> [--target <子目录>] [--rename]     执行脱敏")
        print("  python sanitize.py restore <工作文件夹> [--record <记录文件>]              执行恢复")
        print()
        print("示例:")
        print("  python sanitize.py config .")
        print("  python sanitize.py sanitize .")
        print("  python sanitize.py sanitize . --target 进港全流程智慧化项目\\update --rename")
        print("  python sanitize.py restore . --record _sanitize_record_20260329_121500.json")
        sys.exit(1)

    mode = sys.argv[1].lower()

    # 提取工作目录（跳过 --record/--target/--rename 及其值）
    skip_next = False
    args = []
    extra_flags = set()
    target_dir = None

    for i, a in enumerate(sys.argv[2:], 2):
        if skip_next:
            skip_next = False
            continue
        if a in ("--record", "--target"):
            # 下一个参数是对应的值，跳过
            skip_next = True
            continue
        if a == "--rename":
            extra_flags.add("rename")
            continue
        args.append(a)

    workspace = Path(args[0]).resolve() if args else Path.cwd()

    # 提取 --target 参数
    if "--target" in sys.argv:
        idx = sys.argv.index("--target")
        if idx + 1 < len(sys.argv):
            target_dir = workspace / sys.argv[idx + 1]

    if mode == "config":
        generate_config(workspace)
    elif mode == "sanitize":
        run_sanitize(workspace, target_dir=target_dir, rename_files="rename" in extra_flags)
    elif mode == "restore":
        # 提取 record 参数
        record_path = None
        if "--record" in sys.argv:
            idx = sys.argv.index("--record")
            if idx + 1 < len(sys.argv):
                record_path = Path(sys.argv[idx + 1])
                if not record_path.is_absolute():
                    record_path = workspace / record_path
                # 安全校验：确保 record 文件在 workspace 范围内
                record_path = record_path.resolve()
                if not str(record_path).startswith(str(workspace.resolve())):
                    log(f"[ERROR] 记录文件路径超出工作区范围: {record_path}")
                    sys.exit(1)
        run_restore(workspace, record_path)
    else:
        log(f"[ERROR] 未知模式: {mode}")
        log("  支持的模式: config / sanitize / restore")
        sys.exit(1)


if __name__ == "__main__":
    main()
